import pandas as pd
from docx import Document
import os

# Create docs directory
os.makedirs("documents_test", exist_ok=True)

# 1. Générer la Checklist d'Audit (Excel)
df_checklist = pd.DataFrame({
    "ID_Exigence": ["REQ-001", "REQ-002", "REQ-003"],
    "Description": [
        "La politique qualité est documentée, tenue à jour et communiquée à tout le personnel.",
        "Une revue de direction est planifiée et réalisée au moins une fois par an afin de s'assurer de son adéquation.",
        "Les audits internes sont menés à des intervalles planifiés pour vérifier la conformité du système."
    ],
    "Criticite": ["Haute", "Haute", "Moyenne"]
})
df_checklist.to_excel("documents_test/checklist_audit_mock.xlsx", index=False)
print("Création de checklist_audit_mock.xlsx : OK")

# 2. Générer le Manuel Qualité (Preuve pour REQ-001)
doc1 = Document()
doc1.add_heading("Manuel Qualité QualiNova - v3.0", 0)
doc1.add_paragraph("L'objectif de ce manuel est de définir les lignes directrices de notre système de management.")
doc1.add_heading("1. Politique Qualité", level=1)
doc1.add_paragraph(
    "La direction s'engage formellement au travers de la présente politique qualité à satisfaire "
    "les exigences légales et réglementaires applicables. La politique qualité est tenue à jour "
    "et communiquée à tout le personnel via notre réseau intranet et affichage en salle de pause."
)
doc1.save("documents_test/Manuel_Qualite_v3.docx")
print("Création de Manuel_Qualite_v3.docx : OK")

# 3. Générer la Preuve de Revue de Direction (Preuve pour REQ-002 - Ancienne date pour créer une NC)
doc2 = Document()
doc2.add_heading("Compte Rendu - Revue de Direction 2024", 0)
doc2.add_paragraph("Date de la réunion : 15 Janvier 2024")
doc2.add_paragraph("Participants : PDG, Responsable Qualité, Directeur Opérations.")
doc2.add_heading("Ordre du jour", level=1)
doc2.add_paragraph(
    "- Revue des KPIs de l'année 2023.\n"
    "- Audit interne : résultats satisfaisants, 2 NC mineures clôturées.\n"
    "- Plan de formation 2024 validé."
)
doc2.add_paragraph("PROCHAINE REVUE PRÉVUE EN 2025. AUCUN COMPTE RENDU 2025 N'EST DISPONIBLE POUR LE MOMENT.")
doc2.save("documents_test/CR_Revue_Direction_2024.docx")
print("Création de CR_Revue_Direction_2024.docx : OK")

# 4. Générer la Procédure d'Audit Interne (Preuve pour REQ-003)
doc3 = Document()
doc3.add_heading("Procédure : Audits Internes", 0)
doc3.add_paragraph("Référence : PROC-AUDIT-01")
doc3.add_heading("1. Planification", level=1)
doc3.add_paragraph(
    "Les audits internes QualiNova sont menés à des intervalles planifiés afin de vérifier la "
    "conformité globale du système. Le planning d'audit annuel est établi au plus tard le 15 décembre."
)
doc3.add_heading("2. Réalisation", level=1)
doc3.add_paragraph("Les auditeurs ne doivent pas auditer leur propre département.")
doc3.save("documents_test/Procedure_Audits_Internes.docx")
print("Création de Procedure_Audits_Internes.docx : OK")
